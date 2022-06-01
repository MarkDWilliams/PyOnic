import sys
from PyPDF2 import PdfFileReader, PdfFileWriter
import fitz
import re
from termcolor import colored
import docx


class color:
    PURPLE = '\033[95m'
    CYAN = '\033[96m'
    DARKCYAN = '\033[36m'
    BLUE = '\033[94m'
    GREEN = '\033[92m'
    YELLOW = '\033[93m'
    RED = '\033[91m'
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'
    END = '\033[0m'
    UNBOLD = '\033[0m'


def bold_part(word):
    start, end = split_word(word)
    # bold = color.BOLD+start.strip()
    # unbold = color.UNBOLD+end.strip()
    #
    # return bold+unbold.strip()
    return colored(start,attrs=['bold'])+end

def split_word(word):
    start = word[:len(word)//2]
    end = word[len(word)//2:]
    return (start,end)

def parse_txt(file="sample.txt"):
    #need to have this preserve whitespace
    with open(file,'r') as text:
        paragraph = [line.split() for line in text]

    list1=[]
    list2=[]
    list3=[]
    for sentence in paragraph:
        for word in sentence:
            word = bold_part(word)
            list1.append(word)
        list2.append(list1)
    list3=[' '.join(i) for i in list2]
    print(list3[0])
    with open('output.txt', 'w') as o:
        for ele in list3:
            o.write(ele)
def process_pdf(file_path="sample.pdf"):
    output_doc=docx.Document()
    with fitz.open(file_path) as doc:
        new_text=""
        new_words=[]
        for page in doc:
            if page.number>5:
                #new_text+=process_pdf_page(output_doc,page)
                process_pdf_page_to_doc(output_doc,page)
    output_doc.save(file_path+"-PyOnic.docx")
    print(new_text)

def process_pdf_page(page):

    all_text=page.get_text()
    lines = all_text.split('\n')
    bolded_page =""

    for line in lines:
        bolded_line=""
        words = line.split(' ')
        for word in words:

            bolded = bold_part(word)
            bolded_line+=bolded+" "
        bolded_page=bolded_page+bolded_line+"\n"
    #print(bolded_page)
    return bolded_page
def process_pdf_page_to_doc(doc,page):
    p = doc.add_paragraph()
    all_text=page.get_text()
    lines = all_text.split('\n')
    bolded_page =""
    p = doc.add_paragraph()
    for line in lines:
        bolded_line=""
        words = line.split(' ')
        for word in words:
            start,end =split_word(word)
            p.add_run(start).bold=True
            p.add_run(end+" ")
        #p.add_run("\n")
    #print(bolded_page)
    return bolded_page
def main():
    #parse_txt()
    process_pdf("sample2.pdf")
    return(0)

if __name__ == '__main__':
    main()